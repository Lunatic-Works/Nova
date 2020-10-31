#!/usr/bin/env python3

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.shared import RGBColor
from nova_script_parser import normalize_dialogue, parse_chapters
from strip_code_tex import parse_code, translate

in_filename = 'scenario.txt'
out_filename = 'scenario_no_code.docx'


def main():
    with open(in_filename, 'r', encoding='utf-8') as f:
        chapters = parse_chapters(f)

    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    style = doc.styles['Heading 1']
    style.font.name = 'Microsoft YaHei'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    style = doc.styles.add_style('Character Name', WD_STYLE_TYPE.CHARACTER)
    style.font.name = 'Microsoft YaHei'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    style.font.color.rgb = RGBColor(210, 210, 210)

    style = doc.styles.add_style('BG', WD_STYLE_TYPE.CHARACTER)
    style.font.name = 'Microsoft YaHei'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    style.font.color.rgb = RGBColor(255, 128, 0)

    style = doc.styles.add_style('BGM', WD_STYLE_TYPE.CHARACTER)
    style.font.name = 'Microsoft YaHei'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    style.font.color.rgb = RGBColor(0, 0, 255)

    for chapter_name, entries, _, _ in chapters:
        print(chapter_name)

        doc.add_heading(chapter_name)
        for code, chara_name, dialogue in entries:
            bg_name, bgm_name = parse_code(code, f)
            if bg_name:
                para = doc.add_paragraph()
                para.add_run('场景：' + translate(bg_name), 'BG')
            if bgm_name:
                para = doc.add_paragraph()
                para.add_run('音乐：' + translate(bgm_name), 'BGM')

            dialogue = normalize_dialogue(dialogue, keep_todo=['配音'])
            if dialogue:
                if chara_name:
                    para = doc.add_paragraph()
                    para.add_run(chara_name, 'Character Name')
                    para.add_run(dialogue)
                else:
                    doc.add_paragraph(dialogue)
        doc.add_page_break()

    doc.save(out_filename)


if __name__ == '__main__':
    main()
